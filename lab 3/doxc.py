from docx import Document

# Попробуем добавить код без указания специфичного стиля

# Создаем новый документ
doc = Document()

# Титульный лист
doc.add_heading("МИНОБРНАУКИ РОССИИ", 0)
doc.add_paragraph(
    "федеральное государственное бюджетное образовательное учреждение\n"
    "высшего образования\n"
    "«Санкт-Петербургский государственный морской технический университет» (СПбГМТУ)\n"
    "_______________________________________________________________\n\n"
    "Факультет цифровых промышленных технологий\n"
    "Направление подготовки 09.03.01\n"
    '"Интеллектуальные технологии киберфизических систем"\n\n'
    "Лабораторная работа №3\n"
    "Вариант: Реализация функции XOR с использованием ИНС\n\n\n\n\n"
    "Студент 2 курса группы 20221\n"
    "Очного отделения\n"
    "Руденко Вячеслав Сергеевич\n\n"
    "Проверил:\n"
    "Преподаватель СПбГМТУ\n"
    "Кайнова Татьяна Денисовна\n\n\n\n"
    "Санкт-Петербург\n"
    "2024"
)

# Оглавление
doc.add_heading("Оглавление", level=1)
doc.add_paragraph("1 ОБЩАЯ ЧАСТЬ\t3")
doc.add_paragraph("1.1 Цель работы\t3")
doc.add_paragraph("1.2 Формулировка задачи\t3")
doc.add_paragraph("2 ХОД РАБОТЫ\t4")
doc.add_paragraph("3 ЗАКЛЮЧЕНИЕ\t10")
doc.add_paragraph("4 СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ\t11")
doc.add_paragraph("ПРИЛОЖЕНИЕ А Листинг кода\t12")

# Общая часть
doc.add_heading("1 ОБЩАЯ ЧАСТЬ", level=1)

doc.add_heading("1.1 Цель работы", level=2)
doc.add_paragraph(
    "Реализовать двуслойную нейронную сеть для выполнения операции XOR с использованием "
    "алгоритма обратного распространения ошибки."
)

doc.add_heading("1.2 Формулировка задачи", level=2)
doc.add_paragraph(
    "Создать нейронную сеть с одним скрытым слоем для реализации функции XOR, "
    "используя логистическую сигмоидальную функцию активации и функцию tanh для скрытых нейронов. "
    "Процесс обучения должен включать корректировку весов с использованием обратного распространения ошибки."
)

# Ход работы
doc.add_heading("2 ХОД РАБОТЫ", level=1)

doc.add_paragraph(
    "Для выполнения работы были использованы следующие библиотеки:\n"
    "1. NumPy для работы с массивами и математическими операциями.\n"
)

doc.add_paragraph(
    "Инициализация весов для каждого нейрона производится с помощью функции `neuron_w`, которая задаёт "
    "начальные значения весов случайными числами в диапазоне от -1.0 до 1.0. Функция `forward_pass` "
    "реализует прямой проход, используя функцию активации tanh для скрытых нейронов и логистическую "
    "сигмоиду для выходного нейрона. Функция `backward_pass` рассчитывает ошибку на выходном нейроне "
    "и передаёт её обратно к скрытым нейронам, а функция `adjust_weights` корректирует веса, "
    "используя рассчитанную ошибку и коэффициент обучения."
)

# Код
doc.add_heading("Листинг кода", level=2)
code = """
import numpy as np

np.random.seed(3)
LEARNING_RATE = 0.1
index_list = [0, 1, 2, 3]

x_train = np.array([
    [1.0, -1.0, -1.0],
    [1.0, -1.0, 1.0],
    [1.0, 1.0, -1.0],
    [1.0, 1.0, 1.0]
])

y_train = np.array([0.0, 1.0, 1.0, 0.0])

def neuron_w(input_count):
    weights = np.zeros(input_count + 1)
    for i in range(1, input_count + 1):
        weights[i] = np.random.uniform(-1.0, 1.0)
    return weights

n_w = [neuron_w(2), neuron_w(2), neuron_w(2)]
n_y = [0, 0, 0]
n_error = [0, 0, 0]

def show_learning():
    print('Current weights:')
    for i, w in enumerate(n_w):
        print(f'neuron {i}: w0 = {w[0]:.2f}, w1 = {w[1]:.2f}, w2 = {w[2]:.2f}')
    print('----------------')

def forward_pass(x):
    global n_y
    n_y[0] = np.tanh(np.dot(n_w[0], x))
    n_y[1] = np.tanh(np.dot(n_w[1], x))
    n2_inputs = np.array([1.0, n_y[0], n_y[1]])
    z2 = np.dot(n_w[2], n2_inputs)
    n_y[2] = 1.0 / (1.0 + np.exp(-z2))

def backward_pass(y_truth):
    global n_error
    error_prime = (y_truth - n_y[2])
    derivative = n_y[2] * (1.0 - n_y[2])
    n_error[2] = error_prime * derivative
    derivative = 1.0 - n_y[0]**2
    n_error[0] = n_w[2][1] * n_error[2] * derivative
    derivative = 1.0 - n_y[1]**2
    n_error[1] = n_w[2][2] * n_error[2] * derivative

def adjust_weights(x):
    global n_w
    n_w[0] += (x * LEARNING_RATE * n_error[0])
    n_w[1] += (x * LEARNING_RATE * n_error[1])
    n2_inputs = np.array([1.0, n_y[0], n_y[1]])
    n_w[2] += (n2_inputs * LEARNING_RATE * n_error[2])

all_correct = False
while not all_correct:
    all_correct = True
    np.random.shuffle(index_list)
    for i in index_list:
        forward_pass(x_train[i])
        backward_pass(y_train[i])
        adjust_weights(x_train[i])
        show_learning()
        for i in range(len(x_train)):
            forward_pass(x_train[i])
            print(f'x1 = {x_train[i][1]:.1f}, x2 = {x_train[i][2]:.1f}, y = {n_y[2]:.4f}')
            if (((y_train[i] < 0.5) and (n_y[2] >= 0.5)) or ((y_train[i] >= 0.5) and (n_y[2] < 0.5))):
                all_correct = False
"""
doc.add_paragraph(code)

# Заключение
doc.add_heading("3 ЗАКЛЮЧЕНИЕ", level=1)
doc.add_paragraph(
    "В ходе выполнения лабораторной работы была создана двуслойная нейронная сеть для реализации функции "
    "исключающее ИЛИ (XOR) с использованием алгоритма обратного распространения ошибки. Программа корректно "
    "обучает сеть, отображая изменения весов на каждом шаге."
)

# Список используемых источников
doc.add_heading("4 СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ", level=1)
doc.add_paragraph(
    "ГОСТ Р 7.0.97-2016. Национальный стандарт Российской Федерации. Система стандартов по информации, "
    "библиотечному и издательскому делу. Организационно-распорядительная документация. Требования к "
    "оформлению документов: утвержден и введен в действие Приказом Федерального агентства по техническому "
    "регулированию и метрологии от 14.05.2018 N 244-ст: Дата введения 2018-07-01."
)

# Сохраняем документ
file_path = "Лабораторная работа №3 Руденко Вячеслав.docx"
doc.save(file_path)

file_path
